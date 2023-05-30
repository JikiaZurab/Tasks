from docx import Document
import re
import traceback

def word_replace(match):
    oldWord = match.group(0)
    if oldWord.islower():
        return newWord
    else:
        return newWord.capitalize()


try:
    f = Document('Агрессия.docx')
    newf = Document()

    oldWord, newWord = 'агрессия', 'доброжелательность'
    regEx = re.compile(r'\b' + re.escape(oldWord) + r'|'
            + re.escape(oldWord.capitalize()) + r'\b')

    for par in f.paragraphs:
        newPar = newf.add_paragraph(style = par.style)
        newPar.alignment = par.alignment
        
        for run in par.runs:
            newText = regEx.sub(word_replace, run.text)
            newRun = newPar.add_run(newText)
            newRun.bold = run.bold
            newRun.italic = run.italic
            newRun.underline = run.underline
            newRun.font.size = run.font.size
            newRun.font.name = run.font.name
            newRun.style.name = run.style.name

    newf.save('Доброжелательность.docx')


except Exception as e:
    print('Something went wrong:')
    traceback.print_exc()
